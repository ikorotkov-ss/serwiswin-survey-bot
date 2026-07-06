#!/usr/bin/env python3
"""
Генератор оформленного опросника для SerwisWin в форматах DOCX и PDF
"""

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

OUTPUT_DIR = os.path.expanduser("~/Desktop")

# ============================================================
# DATA
# ============================================================

title = "ОПРОСНИК ДЛЯ КОМАНДЫ SerwisWin"
subtitle = "Сбор инсайтов для рекламных текстов и тем видео-роликов"

sections = [
    {
        "heading": "ВВЕДЕНИЕ",
        "items": [
            {
                "type": "text",
                "content": "Данный опросник предназначен для сбора информации от всех сотрудников,"
                " которые так или иначе контактируют с клиентами SerwisWin."
                " На основе полученных ответов мы будем создавать тексты для таргетированной"
                " рекламы (Facebook/Instagram Ads) и сценарии для видео-роликов.",
            },
            {
                "type": "text",
                "content": "Мы хотим понять реальные боли клиентов, их страхи, возражения, а также то,"
                " что им нравится в нашем сервисе. Нас интересуют живые цитаты и конкретные случаи.",
            },
        ],
    },
    {
        "heading": "ДЛЯ КОГО ОПРОСНИК",
        "items": [
            {
                "type": "list",
                "content": [
                    "Продавцы — те, кто принимает и обрабатывает заявки, обзванивает лидов",
                    "Сотрудники колл-центра — те, кто принимает входящие звонки и консультирует",
                    "Мастера — те, кто выезжает на объекты и выполняет работы на месте",
                ],
            },
        ],
    },
    {
        "heading": "ФОРМАТ СБОРА",
        "items": [
            {
                "type": "text",
                "content": "Ответы собираются в виде голосовых сообщений в Telegram/WhatsApp."
                " Каждый отвечает на вопросы блока, который соответствует его роли.",
            },
            {
                "type": "text",
                "content": "Продолжительность одного голосового ответа — 1-3 минуты."
                " Всего на прохождение своего блока уйдёт около 20-30 минут.",
            },
        ],
    },
    {
        "heading": "ИНСТРУКЦИЯ: КАК ОТВЕЧАТЬ НА ВОПРОСЫ",
        "items": [
            {
                "type": "instruction",
                "steps": [
                    "Читаешь вопрос — вспоминаешь 2-3 конкретных случая из своей практики.",
                    "Отвечаешь голосовым сообщением (1-3 минуты на вопрос).",
                    "Приводи ДОСЛОВНЫЕ цитаты клиентов, если запомнил.",
                    "Если клиент говорил на польском — цитируй по-польски. Если на русском или украинском — соответственно.",
                    "Не приукрашивай — плохой опыт клиента так же важен, как и хороший.",
                    "Говори максимально конкретно. Вместо «клиентам не нравится» → «клиент сказал: \"через щель дует, уже утеплял одеялом, ничего не помогает\"».",
                ],
            },
        ],
    },
    {
        "heading": "ПРИМЕР ХОРОШЕГО ОТВЕТА",
        "items": [
            {
                "type": "example",
                "bold": "Вопрос:",
                "normal": ' "С какой самой частой проблемой приходит клиент?"',
            },
            {
                "type": "example",
                "bold": "Ответ:",
                "normal": ' "На прошлой неделе было три звонка подряд. Первый — пани из Варшавы,'
                ' Жолибож: «Okno się nie zamyka, muszę podpierać kijem, boję się że wypadnie».'
                ' Второй — пан из Мокотова: «Po wymianie szyb w całym bloku wieje, rachunki'
                ' za gaz poszły w górę o 300 zł». Третий — русскоязычная семья из Урсынова:'
                ' «Окна не держат тепло, дети мёрзнут, скажите, это лечится или только менять?»'
                ' Самая частая боль — сквозняки и холод, зимой люди реально мучаются."',
            },
        ],
    },
    {
        "heading": "БЛОК 1. УНИВЕРСАЛЬНЫЙ (ДЛЯ ВСЕХ РОЛЕЙ)",
        "subtitle": "Тема: Типичные клиенты и их проблемы",
        "items": [
            {
                "type": "question",
                "number": 1,
                "text": "Самая частая проблема, с которой к вам обращаются? Назови 3 самые частые.",
            },
            {
                "type": "question",
                "number": 2,
                "text": "В какое время года больше всего обращений? С чем это связано?",
            },
            {
                "type": "question",
                "number": 3,
                "text": "Кто чаще всего обращается? (владельцы квартир/домов, арендаторы, управляющие компаниями, офисы, тип застройки — новостройки или старый фонд)",
            },
            {
                "type": "question",
                "number": 4,
                "text": "Как клиент описывает проблему своими словами? Приведи 2-3 дословные цитаты на польском (и на русском/украинском, если были).",
            },
            {
                "type": "question",
                "number": 5,
                "text": "Какие страхи и опасения высказывают клиенты до того, как соглашаются на услугу? («А не обманут?», «А не сломают ещё больше?», «Это дорого?»)",
            },
            {
                "type": "question",
                "number": 6,
                "text": "Что клиенты говорят, когда отказываются от услуги после консультации или оценки? Какая причина отказа самая частая?",
            },
        ],
    },
    {
        "heading": "БЛОК 1 (продолжение)",
        "subtitle": "Тема: Процесс принятия решения",
        "items": [
            {
                "type": "question",
                "number": 7,
                "text": "Почему клиенты выбирают именно вас? Что они называют в разговоре — цена, скорость, гарантия, отзывы, диагностика, сертификат ISO?",
            },
            {
                "type": "question",
                "number": 8,
                "text": "Что клиенты ценят в сервисе больше всего после выполнения работы? Самые частые комплименты и слова благодарности.",
            },
            {
                "type": "question",
                "number": 9,
                "text": "Бывают ли моменты, когда клиент говорит «вау»? Что конкретно вызывает такую реакцию?",
            },
            {
                "type": "question",
                "number": 10,
                "text": "Какие возражения вы слышите чаще всего и как вы на них отвечаете?",
            },
            {
                "type": "question",
                "number": 11,
                "text": "Как клиенты реагируют на гарантию 24 месяца? Упоминаете ли вы её сами? Влияет ли это на решение?",
            },
            {
                "type": "question",
                "number": 12,
                "text": "С чем клиенты сравнивают цены? (с новыми окнами, с другой фирмой, «сам сделаю», «сосед дешевле сделал»)",
            },
            {
                "type": "question",
                "number": 13,
                "text": "Как клиенты реагируют на срочность — «90% устраняем в день обращения»? Это влияет на выбор? Приводят ли примеры других фирм, где ждали неделями?",
            },
            {
                "type": "question",
                "number": 14,
                "text": "Что клиентам НЕ нравится в нашем сервисе? Бывают ли жалобы? (долгое ожидание, цена, неудобное время, недоделки, вежливость сотрудников)",
            },
            {
                "type": "question",
                "number": 15,
                "text": "Есть ли разница между польскоязычными и русско/украиноязычными клиентами? Разные боли? Разные возражения? Кто легче соглашается?",
            },
        ],
    },
    {
        "heading": "БЛОК 2. ДЛЯ ПРОДАВЦОВ / КОЛЛ-ЦЕНТРА",
        "subtitle": "Кто общается с клиентом до выезда мастера",
        "items": [
            {
                "type": "question",
                "number": 16,
                "text": "Как проходит типичный первый звонок? Что клиент говорит в первые 30 секунд?",
            },
            {
                "type": "question",
                "number": 17,
                "text": "Какие вопросы клиенты задают чаще всего до того, как согласиться на диагностику или выезд? («Сколько стоит?», «Как долго?», «Есть ли гарантия?»)",
            },
            {
                "type": "question",
                "number": 18,
                "text": "Что клиенты НЕ говорят, но вы чувствуете или понимаете? (скрытые сомнения, невысказанные страхи)",
            },
            {
                "type": "question",
                "number": 19,
                "text": "Какая информация, полученная от клиента, помогает вам понять — это «горячий» лид? По каким фразам вы понимаете, что клиент готов заказать прямо сейчас?",
            },
            {
                "type": "question",
                "number": 20,
                "text": "Какие типы клиентов самые сложные? (недоверчивые, торгующиеся, не понимающие, чего хотят, требующие скидку)",
            },
            {
                "type": "question",
                "number": 21,
                "text": "Как клиенты реагируют на цену? Какая реакция на разные услуги: регулировка от 65 zł, замена стекла от 800 zł, замена механизма от 1210 zł?",
            },
            {
                "type": "question",
                "number": 22,
                "text": "Какие вопросы вы задаёте клиенту, чтобы понять, что именно нужно? Скрипт или ключевые вопросы.",
            },
        ],
    },
    {
        "heading": "БЛОК 2 (продолжение)",
        "subtitle": "Тема: Бесплатная диагностика как триггер",
        "items": [
            {
                "type": "question",
                "number": 23,
                "text": "Как клиенты реагируют на предложение бесплатной диагностики? Что говорят?",
            },
            {
                "type": "question",
                "number": 24,
                "text": "Часто ли клиенты соглашаются на диагностику, но потом не делают ремонт? Что они говорят в таком случае?",
            },
            {
                "type": "question",
                "number": 25,
                "text": "Что клиенты говорят про «гарантию качества — ISO 9001»? Влияет ли это на решение?",
            },
        ],
    },
    {
        "heading": "БЛОК 2B. РЕНОВАЦИЯ ДЕРЕВЯННЫХ ОКОН (ДЛЯ МАСТЕРОВ)",
        "subtitle": "Отдельный блок для мастеров, работающих с деревянными окнами",
        "items": [
            {
                "type": "instruction",
                "steps": [
                    "Этот блок только для тех мастеров, кто делал реставрацию/реновацию деревянных окон (шлифовка, пропитка, покраса). Если не делали — пропустите.",
                ],
            },
            {
                "type": "question",
                "number": 26,
                "text": "Почему клиент решается на реновацию, а не на замену окон? Какие аргументы приводит? («дешевле», «окна старые, но качественные», «не хочу менять рамы»)?",
            },
            {
                "type": "question",
                "number": 27,
                "text": "Как клиент реагирует на результат реновации? Сравнивает с новыми окнами? Какие эмоции — восторг, удивление, благодарность?",
            },
            {
                "type": "question",
                "number": 28,
                "text": "Какие проблемы деревянных окон вы чаще всего видите? (гниль, рассохлись, краска облупилась, не закрываются, щели между рамой и стеной)",
            },
            {
                "type": "question",
                "number": 29,
                "text": "Сколько времени обычно занимает реновация одного окна? Как клиент реагирует на то, что это не один день?",
            },
            {
                "type": "question",
                "number": 30,
                "text": "Что клиенты говорят про запах краски/лака? Мешает ли это, просят ли делать в определённое время?",
            },
            {
                "type": "question",
                "number": 31,
                "text": "Какие районы или дома чаще заказывают реновацию? Старый фонд, довоенные дома, villas, конкретные районы Варшавы (Śródmieście, Żoliborz, Mokotów)?",
            },
        ],
    },
    {
        "heading": "БЛОК 3. ДЛЯ МАСТЕРОВ (ВЫЕЗДНЫХ)",
        "subtitle": "Кто работает на объекте",
        "items": [
            {
                "type": "question",
                "number": 32,
                "text": "Что клиент говорит или показывает, когда вы приезжаете? Какая первая фраза или реакция?",
            },
            {
                "type": "question",
                "number": 33,
                "text": "Какую реакцию клиента вы замечаете, когда показываете проблему с помощью камеры/анемометра? Что конкретно говорят?",
            },
            {
                "type": "question",
                "number": 34,
                "text": "Что клиенты больше всего ценят в вашей работе непосредственно в процессе? (чистота, аккуратность, скорость, объяснения)",
            },
            {
                "type": "question",
                "number": 35,
                "text": "Какие фразы или комментарии вы слышите от клиента, когда работа уже сделана? Самые яркие моменты.",
            },
            {
                "type": "question",
                "number": 36,
                "text": "Бывали ли случаи, когда клиент не ожидал такого результата? Расскажи 1-2 истории.",
            },
            {
                "type": "question",
                "number": 37,
                "text": "Что клиенты говорят про окна после ремонта? Сравнивают с новыми? Говорят, что стало теплее или тише?",
            },
        ],
    },
    {
        "heading": "БЛОК 3 (продолжение)",
        "subtitle": "Тема: Типичные наблюдения",
        "items": [
            {
                "type": "question",
                "number": 38,
                "text": "С какой неожиданной проблемой вы сталкивались на месте, которую клиент не описал по телефону? Примеры.",
            },
            {
                "type": "question",
                "number": 39,
                "text": "Как часто проблема вызвана неправильной эксплуатацией, а не дефектом окон? Что говорят клиенты, когда вы это объясняете?",
            },
            {
                "type": "question",
                "number": 40,
                "text": "Какие районы или типы застройки чаще всего имеют одинаковые проблемы? (новостройки, старый фонд, конкретные районы Варшавы)",
            },
            {
                "type": "question",
                "number": 41,
                "text": "Что вы слышите от клиентов про наш сервисный отчёт (certyfikowane sprawozdanie techniczne)? Читают? Ценят? Игнорируют?",
            },
        ],
    },
    {
        "heading": "ФИНАЛЬНЫЙ БЛОК (ДЛЯ ВСЕХ)",
        "items": [
            {
                "type": "question",
                "number": 42,
                "text": "Какую рекламу или объявление вы бы сделали сами, если бы были маркетологом? Какую фразу или историю поставили бы в центр рекламы?",
            },
            {
                "type": "question",
                "number": 43,
                "text": "Если одним предложением — почему клиент должен выбрать SerwisWin вместо альтернатив? Твоя версия.",
            },
            {
                "type": "question",
                "number": 44,
                "text": "Что бы вы изменили в сервисе, чтобы клиентам стало ещё приятнее или удобнее?",
            },
            {
                "type": "question",
                "number": 45,
                "text": "Какие вопросы мы забыли? Добавь то, что считаешь важным.",
            },
        ],
    },
    {
        "heading": "КАК ИСПОЛЬЗОВАТЬ РЕЗУЛЬТАТЫ",
        "items": [
            {
                "type": "text",
                "content": "После сбора голосовых ответов каждый ответ и цитату нужно отнести к одной из категорий:",
            },
            {
                "type": "list",
                "content": [
                    "БОЛЬ — холод, сквозняк, шум, поломка, безопасность, высокие счета за отопление",
                    "СТРАХ — обман, дорого, ещё хуже сломают",
                    "РАДОСТЬ — тепло, тихо, чисто, быстро, дешевле чем новые окна",
                    "ВОЗРАЖЕНИЕ — дорого, подумаю, сравню",
                    "ВОСХИЩЕНИЕ — как новые, спасибо, не ожидал",
                ],
            },
        ],
    },
    {
        "heading": "ЧТО МЫ ПОЛУЧАЕМ НА ВЫХОДЕ",
        "items": [
            {
                "type": "numbered",
                "content": [
                    {
                        "bold": "Тексты для таргета (Facebook / Instagram Ads):",
                        "normal": " начинаем с прямой боли или цитаты клиента.",
                    },
                    {
                        "bold": "Пример:",
                        "normal": ' «Okno się nie domyka? Naprawimy je jeszcze dziś. Darmowy dojazd + diagnoza.»',
                    },
                    {
                        "bold": "Пример:",
                        "normal": ' «Через окна дует? Счета за газ растут? Приедем, найдём проблему тепловизором — бесплатно.»',
                    },
                ],
            },
            {
                "type": "numbered",
                "content": [
                    {
                        "bold": "Темы для видео-роликов:",
                        "normal": " разбор конкретных проблем с реальными историями.",
                    },
                    {
                        "bold": "Пример:",
                        "normal": ' «Пани Мария думала, что нужно менять все окна. Мы отрегулировали и заменили уплотнители — стало тепло, и она сэкономила 15 000 zł.»',
                    },
                    {
                        "bold": "Пример:",
                        "normal": ' «3 признака, что ваши окна пора регулировать. Проверьте себя за 2 минуты.»',
                    },
                ],
            },
            {
                "type": "numbered",
                "content": [
                    {
                        "bold": "A/B-тесты формулировок:",
                        "normal": " сравниваем эффективность формулировок через боль vs выгоду vs страх.",
                    },
                ],
            },
        ],
    },
]


# ============================================================
# HELPERS
# ============================================================

def set_cell_shading(cell, color):
    shading_elm = OxmlElement("w:shd")
    shading_elm.set(qn("w:fill"), color)
    shading_elm.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(shading_elm)


def add_styled_paragraph(doc, text, style="Normal", bold=False, size=None, color=None, alignment=None, space_after=None, space_before=None, italic=False):
    p = doc.add_paragraph(style=style)
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    if size:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*color)
    if alignment is not None:
        p.alignment = alignment
    if space_after is not None:
        p.paragraph_format.space_after = Pt(space_after)
    if space_before is not None:
        p.paragraph_format.space_before = Pt(space_before)
    return p, run


def add_rich_paragraph(doc, parts, alignment=None, space_after=None, space_before=None):
    """parts = list of dicts with 'text', 'bold', 'italic', 'size', 'color' (optional)"""
    p = doc.add_paragraph()
    if alignment:
        p.alignment = alignment
    if space_after is not None:
        p.paragraph_format.space_after = Pt(space_after)
    if space_before is not None:
        p.paragraph_format.space_before = Pt(space_before)
    for part in parts:
        run = p.add_run(part.get("text", ""))
        run.bold = part.get("bold", False)
        run.italic = part.get("italic", False)
        if "size" in part:
            run.font.size = Pt(part["size"])
        if "color" in part:
            run.font.color.rgb = RGBColor(*part["color"])
    return p


def build_document():
    doc = Document()

    # ---- Page setup ----
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    # ---- Styles ----
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)
    font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    # ---- TITLE ----
    add_styled_paragraph(
        doc, "ОПРОСНИК ДЛЯ КОМАНДЫ SerwisWin",
        bold=True, size=24, color=(0x1A, 0x56, 0x8E),
        alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=4
    )
    add_styled_paragraph(
        doc, "Сбор инсайтов для рекламных текстов и тем видео-роликов",
        bold=False, size=14, color=(0x66, 0x66, 0x66),
        alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=6
    )

    # Thin line
    p_line = doc.add_paragraph()
    p_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_line = p_line.add_run("━" * 60)
    run_line.font.color.rgb = RGBColor(0x1A, 0x56, 0x8E)
    run_line.font.size = Pt(6)
    p_line.paragraph_format.space_after = Pt(8)

    # ---- SECTIONS ----
    for sec in sections:
        heading = sec["heading"]

        # Section heading
        h = doc.add_heading(heading, level=1)
        for run in h.runs:
            run.font.color.rgb = RGBColor(0x1A, 0x56, 0x8E)
            run.font.size = Pt(16)
            run.font.name = "Calibri"
        h.paragraph_format.space_before = Pt(18)
        h.paragraph_format.space_after = Pt(6)

        # Subtitle
        if "subtitle" in sec:
            add_styled_paragraph(
                doc, sec["subtitle"],
                bold=True, italic=True, size=11,
                color=(0x55, 0x55, 0x55), space_after=8
            )

        for item in sec["items"]:
            t = item["type"]

            if t == "text":
                p = doc.add_paragraph(item["content"])
                p.paragraph_format.space_after = Pt(6)
                p.paragraph_format.line_spacing = 1.15

            elif t == "list":
                for li in item["content"]:
                    p = doc.add_paragraph(style="List Bullet")
                    p.text = li
                    p.paragraph_format.space_after = Pt(2)
                    p.paragraph_format.space_before = Pt(1)
                    p.paragraph_format.left_indent = Cm(0.8)

            elif t == "instruction":
                for i, step in enumerate(item["steps"], 1):
                    p = doc.add_paragraph()
                    p.paragraph_format.space_after = Pt(3)
                    p.paragraph_format.left_indent = Cm(0.8)
                    run_num = p.add_run(f"{i}. ")
                    run_num.bold = True
                    run_num.font.size = Pt(11)
                    run_num.font.color.rgb = RGBColor(0x1A, 0x56, 0x8E)
                    run_text = p.add_run(step)
                    run_text.font.size = Pt(11)

            elif t == "example":
                p = doc.add_paragraph()
                p.paragraph_format.space_after = Pt(3)
                p.paragraph_format.left_indent = Cm(0.8)
                run_b = p.add_run(item.get("bold", ""))
                run_b.bold = True
                run_b.font.size = Pt(11)
                run_n = p.add_run(item.get("normal", ""))
                run_n.font.size = Pt(11)
                run_n.italic = True
                # light gray background in example — via table (not great for pdf but fine)
                p.paragraph_format.space_after = Pt(2)

            elif t == "question":
                p = doc.add_paragraph()
                p.paragraph_format.space_after = Pt(6)
                p.paragraph_format.left_indent = Cm(0.6)
                p.paragraph_format.space_before = Pt(4)

                run_num = p.add_run(f"{item['number']}. ")
                run_num.bold = True
                run_num.font.size = Pt(11)
                run_num.font.color.rgb = RGBColor(0x1A, 0x56, 0x8E)

                # Check if there's a question number assigned vs title
                qtext = item.get("text", "")
                run_text = p.add_run(qtext)
                run_text.font.size = Pt(11)

            elif t == "numbered":
                for i, sub in enumerate(item.get("content", []), 1):
                    p = doc.add_paragraph()
                    p.paragraph_format.space_after = Pt(4)
                    p.paragraph_format.left_indent = Cm(1.0)
                    run_n = p.add_run(f"{i}. ")
                    run_n.bold = True
                    run_n.font.color.rgb = RGBColor(0x1A, 0x56, 0x8E)
                    run_n.font.size = Pt(11)
                    if "bold" in sub:
                        run_b = p.add_run(sub["bold"] + " ")
                        run_b.bold = True
                        run_b.font.size = Pt(11)
                    if "normal" in sub:
                        run_t = p.add_run(sub["normal"])
                        run_t.font.size = Pt(11)

        # Extra spacer
        doc.add_paragraph().paragraph_format.space_after = Pt(2)

    # ---- FOOTER ----
    add_styled_paragraph(
        doc, "━" * 60,
        size=6, color=(0x1A, 0x56, 0x8E),
        alignment=WD_ALIGN_PARAGRAPH.CENTER, space_before=12, space_after=2
    )
    add_styled_paragraph(
        doc, "SerwisWin — Naprawa okien PCV i drewnianych | serwiswin.pl | +48 579 779 611",
        size=9, color=(0x99, 0x99, 0x99),
        alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=2
    )

    return doc


def generate_pdf(pdf_path):
    """Generate a professionally designed PDF with proper layout, colors, and typography."""
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.lib.units import cm, mm
    from reportlab.lib.colors import HexColor, white, black
    from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT, TA_JUSTIFY
    from reportlab.platypus import (
        SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle,
        PageBreak, ListFlowable, ListItem, KeepTogether
    )
    from reportlab.platypus.flowables import HRFlowable
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont

    # ---------- Fonts ----------
    font_path = "/Library/Fonts/Arial Unicode.ttf"
    pdfmetrics.registerFont(TTFont("ArialUnicode", font_path))
    # ReportLab only uses one font — we map all variants to the same TTF (Arial Unicode has all weights)

    F = "ArialUnicode"

    # ---------- Colour palette ----------
    BLUE_DARK   = HexColor("#0D3B66")   # deep navy
    BLUE_MID    = HexColor("#1A568E")   # primary blue
    BLUE_LIGHT  = HexColor("#E8F0FE")   # very light blue bg
    BLUE_ACCENT = HexColor("#B3D4FC")   # medium light for borders
    GRAY_DARK   = HexColor("#333333")
    GRAY_MID    = HexColor("#666666")
    GRAY_LIGHT  = HexColor("#BBBBBB")
    OFF_WHITE   = HexColor("#FAFBFC")
    WHITE       = white
    ORANGE_ACC  = HexColor("#E67E22")   # accent for highlights

    # ---------- Styles ----------
    s_title = ParagraphStyle("s_title", fontName=F, fontSize=22, leading=26,
                              textColor=WHITE, alignment=TA_CENTER, spaceAfter=2)
    s_subtitle = ParagraphStyle("s_subtitle", fontName=F, fontSize=11, leading=14,
                                 textColor=HexColor("#D0E0F0"), alignment=TA_CENTER, spaceAfter=2)
    s_tagline = ParagraphStyle("s_tagline", fontName=F, fontSize=9, leading=11,
                                textColor=HexColor("#A0B8D0"), alignment=TA_CENTER)

    s_h1 = ParagraphStyle("s_h1", fontName=F, fontSize=13, leading=16,
                           textColor=WHITE, alignment=TA_LEFT, spaceBefore=0, spaceAfter=0)
    s_h2_sub = ParagraphStyle("s_h2_sub", fontName=F, fontSize=10, leading=13,
                               textColor=GRAY_MID, spaceAfter=6, leftIndent=2)
    s_body = ParagraphStyle("s_body", fontName=F, fontSize=9.5, leading=13,
                             textColor=GRAY_DARK, spaceAfter=5, alignment=TA_JUSTIFY)
    s_bullet = ParagraphStyle("s_bullet", fontName=F, fontSize=9.5, leading=13,
                               textColor=GRAY_DARK, spaceAfter=3, leftIndent=22,
                               bulletIndent=12)
    s_instr = ParagraphStyle("s_instr", fontName=F, fontSize=9.5, leading=13,
                              textColor=GRAY_DARK, spaceAfter=3, leftIndent=14)
    s_example = ParagraphStyle("s_example", fontName=F, fontSize=9.5, leading=13,
                                textColor=GRAY_DARK, spaceAfter=4, leftIndent=14)
    s_question = ParagraphStyle("s_question", fontName=F, fontSize=9.5, leading=13.5,
                                 textColor=GRAY_DARK, spaceAfter=6, leftIndent=6, spaceBefore=2)
    s_num_item = ParagraphStyle("s_num_item", fontName=F, fontSize=9.5, leading=13,
                                 textColor=GRAY_DARK, spaceAfter=4, leftIndent=18)
    s_footer = ParagraphStyle("s_footer", fontName=F, fontSize=7.5, leading=9.5,
                               textColor=GRAY_LIGHT, alignment=TA_CENTER)
    s_page_num = ParagraphStyle("s_page_num", fontName=F, fontSize=7.5, leading=9.5,
                                 textColor=GRAY_MID, alignment=TA_CENTER)

    # ---------- Helpers ----------

    def PN(text, style=s_body):
        """Clean paragraph — auto-escape XML special chars."""
        return Paragraph(
            text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
            .replace("(C)", "&copy;").replace("(R)", "&reg;"),
            style
        )

    def PB(text, style=s_body):
        """Paragraph that allows inline HTML tags (font, b, i, etc.)."""
        return Paragraph(text, style)

    def colored_bullet(text, style=s_bullet):
        """Bullet with a blue • instead of &bull; entity."""
        return Paragraph(
            f'<font color="{BLUE_MID.hexval()}">•</font>&nbsp;&nbsp;{text}',
            style
        )

    def make_section_header(title):
        """Return a full-width blue bar with white title text."""
        bar_data = [[Paragraph(f"<b>{title}</b>", s_h1)]]
        bar = Table(bar_data, colWidths=[17 * cm])
        bar.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, -1), BLUE_MID),
            ("LEFTPADDING", (0, 0), (-1, -1), 14),
            ("RIGHTPADDING", (0, 0), (-1, -1), 14),
            ("TOPPADDING", (0, 0), (-1, -1), 7),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 7),
            ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ]))
        return bar

    def make_question_card(number, text):
        """Return a light-blue card containing a numbered question."""
        q = Paragraph(
            f'<font color="{BLUE_MID.hexval()}"><b>{number}.</b></font>&nbsp;&nbsp;{text}',
            s_question
        )
        card = Table([[q]], colWidths=[16.2 * cm])
        card.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, -1), BLUE_LIGHT),
            ("LEFTPADDING", (0, 0), (-1, -1), 10),
            ("RIGHTPADDING", (0, 0), (-1, -1), 10),
            ("TOPPADDING", (0, 0), (-1, -1), 5),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
            ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
            ("BOX", (0, 0), (-1, -1), 0.4, BLUE_ACCENT),
        ]))
        return card

    def make_instruction_card(steps):
        """Return a bordered card with numbered instruction steps."""
        rows = []
        for i, step in enumerate(steps, 1):
            p = Paragraph(
                f'<font color="{BLUE_MID.hexval()}"><b>{i}.</b></font> {step}',
                s_instr
            )
            rows.append([p])
        card = Table(rows, colWidths=[16.2 * cm])
        style_cmds = [
            ("BACKGROUND", (0, 0), (-1, -1), HexColor("#FFF8E7")),
            ("LEFTPADDING", (0, 0), (-1, -1), 12),
            ("RIGHTPADDING", (0, 0), (-1, -1), 12),
            ("TOPPADDING", (0, 0), (-1, -1), 5),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
            ("BOX", (0, 0), (-1, -1), 0.5, ORANGE_ACC),
            ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ]
        for i in range(len(steps) - 1):
            style_cmds.append(("LINEBELOW", (0, i), (-1, i), 0.3, HexColor("#F0E0C0")))
        card.setStyle(TableStyle(style_cmds))
        return card

    def make_example_card(bold_text, normal_text):
        """Return a card for the example block with subtle background."""
        p = Paragraph(f"<b>{bold_text}</b><i>{normal_text}</i>", s_example)
        card = Table([[p]], colWidths=[16.2 * cm])
        card.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, -1), HexColor("#F5F5F5")),
            ("LEFTPADDING", (0, 0), (-1, -1), 12),
            ("RIGHTPADDING", (0, 0), (-1, -1), 12),
            ("TOPPADDING", (0, 0), (-1, -1), 6),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
            ("BOX", (0, 0), (-1, -1), 0.3, GRAY_LIGHT),
        ]))
        return card

    # ---------- Flowables ----------

    flowables = []

    # ========================================
    # COVER / TITLE BLOCK
    # ========================================
    # Full-width dark navy header box
    cover_data = [
        [Paragraph("<b>ОПРОСНИК ДЛЯ КОМАНДЫ</b>", s_title)],
        [Paragraph("<b>SerwisWin</b>", ParagraphStyle("s_title_big", fontName=F, fontSize=26, leading=30,
                                                       textColor=HexColor("#FFD700"), alignment=TA_CENTER,
                                                       spaceBefore=4, spaceAfter=6))],
        [Paragraph("Сбор инсайтов для рекламных текстов и тем видео-роликов", s_subtitle)],
        [Spacer(1, 6)],
        [Paragraph("Сотрудники, работающие с клиентами &mdash; "
                   "продавцы, колл-центр, выездные мастера",
                   s_tagline)],
        [Spacer(1, 16)],
        [Paragraph(f'<font color="{BLUE_MID.hexval()}">─</font>' * 50,
                   ParagraphStyle("dashline", fontName=F, fontSize=6, leading=7,
                                  textColor=HexColor("#4A7DB5"), alignment=TA_CENTER,
                                  spaceBefore=2))],
    ]
    cover_table = Table(cover_data, colWidths=[17 * cm])
    cover_table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, -1), BLUE_DARK),
        ("LEFTPADDING", (0, 0), (-1, -1), 20),
        ("RIGHTPADDING", (0, 0), (-1, -1), 20),
        ("TOPPADDING", (0, 0), (-1, -1), 30),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 25),
        ("ALIGN", (0, 0), (-1, -1), "CENTER"),
    ]))
    flowables.append(cover_table)
    flowables.append(Spacer(1, 10))

    # Thin accent line below cover
    flowables.append(HRFlowable(width="100%", thickness=1.5, color=BLUE_MID,
                                 spaceAfter=6, spaceBefore=0))

    # Branding bar
    brand_data = [[Paragraph(
        'SerwisWin.pl &nbsp;|&nbsp; +48 579 779 611 &nbsp;|&nbsp; '
        'Naprawa okien PCV i drewnianych &nbsp;|&nbsp; '
        '&#9733; 4.9 Google (1929 opinii) &nbsp;|&nbsp; ISO 9001',
        ParagraphStyle("brandbar", fontName=F, fontSize=7, leading=9,
                       textColor=WHITE, alignment=TA_CENTER))]]
    brand_table = Table(brand_data, colWidths=[17 * cm])
    brand_table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, -1), BLUE_MID),
        ("LEFTPADDING", (0, 0), (-1, -1), 8),
        ("RIGHTPADDING", (0, 0), (-1, -1), 8),
        ("TOPPADDING", (0, 0), (-1, -1), 4),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
    ]))
    flowables.append(brand_table)
    flowables.append(Spacer(1, 12))

    # ========================================
    # SECTIONS
    # ========================================
    for sec in sections:
        heading = sec["heading"]
        subtitle = sec.get("subtitle")

        # Section header bar
        flowables.append(make_section_header(heading))
        flowables.append(Spacer(1, 4))

        # Subtitle line
        if subtitle:
            flowables.append(PB(f'<i>{subtitle}</i>', s_h2_sub))
            flowables.append(Spacer(1, 2))

        for item in sec["items"]:
            t = item["type"]

            if t == "text":
                flowables.append(PN(item["content"], s_body))

            elif t == "list":
                for li in item["content"]:
                    flowables.append(colored_bullet(li))

            elif t == "instruction":
                flowables.append(make_instruction_card(item["steps"]))
                flowables.append(Spacer(1, 4))

            elif t == "example":
                flowables.append(make_example_card(item.get("bold", ""), item.get("normal", "")))
                flowables.append(Spacer(1, 4))

            elif t == "question":
                flowables.append(make_question_card(item["number"], item.get("text", "")))
                flowables.append(Spacer(1, 3))

            elif t == "numbered":
                for i, sub in enumerate(item.get("content", []), 1):
                    b = sub.get("bold", "")
                    n = sub.get("normal", "")
                    flowables.append(PB(
                        f'<font color="{BLUE_MID.hexval()}"><b>{i}.</b></font> '
                        f'<b>{b}</b> {n}',
                        s_num_item
                    ))

        # Spacer between sections
        flowables.append(Spacer(1, 8))

    # ========================================
    # FOOTER
    # ========================================
    flowables.append(HRFlowable(width="100%", thickness=1, color=BLUE_MID, spaceAfter=4, spaceBefore=8))
    flowables.append(PB(
        'SerwisWin &mdash; Naprawa okien PCV i drewnianych<br/>'
        'serwiswin.pl &nbsp;|&nbsp; +48 579 779 611 &nbsp;|&nbsp; '
        'NIP 5223253554 &nbsp;|&nbsp; Regon 388056300',
        s_footer
    ))
    flowables.append(Spacer(1, 6))

    # ========================================
    # PAGE NUMBER (canvas callback)
    # ========================================
    def add_page_number(canvas, doc):
        canvas.saveState()
        canvas.setFont(F, 7.5)
        canvas.setFillColor(GRAY_MID)
        canvas.drawCentredString(A4[0] / 2.0, 1.2 * cm,
                                 f"— {doc.page} —")
        canvas.restoreState()

    # ========================================
    # BUILD
    # ========================================
    doc_template = SimpleDocTemplate(
        pdf_path,
        pagesize=A4,
        rightMargin=1.8 * cm,
        leftMargin=1.8 * cm,
        topMargin=1.5 * cm,
        bottomMargin=1.8 * cm,
    )
    doc_template.build(flowables, onFirstPage=add_page_number, onLaterPages=add_page_number)


# ============================================================
# MAIN
# ============================================================

if __name__ == "__main__":
    docx_path = os.path.join(OUTPUT_DIR, "SerwisWin_Опросник.docx")
    pdf_path = os.path.join(OUTPUT_DIR, "SerwisWin_Опросник.pdf")

    print("▶ Generating DOCX...")
    doc = build_document()
    doc.save(docx_path)
    print(f"  ✓ DOCX saved → {docx_path}")

    print("▶ Generating PDF (with Cyrillic font)...")
    generate_pdf(pdf_path)
    print(f"  ✓ PDF saved → {pdf_path}")

    print("\n✅ Done! Both files on your Desktop.")
